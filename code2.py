import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import logging
from io import BytesIO, StringIO
import datetime

# Set up logging and environment variables
logging.basicConfig(level=logging.INFO)
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    raise ValueError("Google API Key not found in environment variables")

# Configure Google Generative AI API
import google.generativeai as genai
genai.configure(api_key=api_key)

def read_pdf(file):
    """Extract text from PDF file."""
    try:
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        logging.error(f"Error processing PDF file: {str(e)}")
        st.error(f"Error processing {file.name}: {str(e)}")
        return ""

def read_docx(file):
    """Extract text from DOCX file."""
    try:
        doc = Document(BytesIO(file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        logging.error(f"Error processing DOCX file: {str(e)}")
        st.error(f"Error processing {file.name}: {str(e)}")
        return ""

def read_txt(file):
    """Extract text from TXT file."""
    try:
        content = file.read()
        # Handle both string and bytes content
        if isinstance(content, bytes):
            return content.decode('utf-8')
        return content
    except Exception as e:
        logging.error(f"Error processing TXT file: {str(e)}")
        st.error(f"Error processing {file.name}: {str(e)}")
        return ""

def read_csv(file):
    """Extract text from CSV file."""
    try:
        df = pd.read_csv(file)
        return df.to_string()
    except Exception as e:
        logging.error(f"Error processing CSV file: {str(e)}")
        st.error(f"Error processing {file.name}: {str(e)}")
        return ""

def read_excel(file):
    """Extract text from Excel file."""
    try:
        df = pd.read_excel(file)
        return df.to_string()
    except Exception as e:
        logging.error(f"Error processing Excel file: {str(e)}")
        st.error(f"Error processing {file.name}: {str(e)}")
        return ""

def get_document_text(docs):
    """Extract text from multiple file formats."""
    text = ""
    for doc in docs:
        try:
            file_extension = doc.name.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                text += read_pdf(doc) + "\n\n"
            elif file_extension == 'docx':
                text += read_docx(doc) + "\n\n"
            elif file_extension == 'txt':
                text += read_txt(doc) + "\n\n"
            elif file_extension == 'csv':
                text += read_csv(doc) + "\n\n"
            elif file_extension in ['xlsx', 'xls']:
                text += read_excel(doc) + "\n\n"
            else:
                st.warning(f"Unsupported file format: {file_extension}")
                continue
                
            logging.info(f"Successfully processed {doc.name}")
        except Exception as e:
            st.error(f"Error processing {doc.name}: {str(e)}")
            logging.error(f"Error processing {doc.name}: {str(e)}")
            continue
            
    return text

def get_text_chunks(text, chunk_size=10000, chunk_overlap=1000):
    """Split text into manageable chunks."""
    try:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        chunks = text_splitter.split_text(text)
        if not chunks:
            raise ValueError("Text chunks are empty.")
        logging.info(f"Text chunks created successfully.")
        return chunks
    except Exception as e:
        logging.error(f"Error creating text chunks: {str(e)}")
        st.error(f"Error creating text chunks: {str(e)}")
        return []

def get_vector_store(text_chunks):
    """Generate or load a FAISS vector store from text chunks."""
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
        vector_store.save_local("faiss_index")
        logging.info("FAISS index created and saved.")
        return vector_store
    except Exception as e:
        logging.error(f"Error creating FAISS vector store: {str(e)}")
        st.error(f"Error creating FAISS vector store: {str(e)}")
        return None

def get_conversational_chain(temperature=0.3):
    """Create a conversational chain using a custom prompt template."""
    try:
        prompt_template = """
        Answer the question as detailed as possible from the provided context. If the answer is not in the context, respond with 'answer is not available in the context'.
        
        Context:\n{context}\n
        Question: \n{question}\n
        Answer:
        """
        model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=temperature)
        prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
        chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
        return chain
    except Exception as e:
        logging.error(f"Error creating conversational chain: {str(e)}")
        st.error(f"Error creating conversational chain: {str(e)}")
        return None

def user_input(user_question, vector_store, chain):
    """Process user question and retrieve answer using the FAISS vector store."""
    try:
        if vector_store is None:
            raise ValueError("No vector store available. Please upload and process documents first.")
        
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        docs = new_db.similarity_search(user_question)
        response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
        return response["output_text"]
    except Exception as e:
        logging.error(f"Error processing user question: {str(e)}")
        st.error(f"Error processing your question: {str(e)}")
        return "An error occurred while processing your question. Please try again."

def main():
    """Main Streamlit app function."""
    st.set_page_config(page_title="Chat with Documents", layout="wide")
    st.title("🗂️ Multi-Format Document Chat Using Gemini")
    st.markdown(
        """
        **Welcome to the Document Chat Interface!**  
        Upload your documents (PDF, DOCX, TXT, CSV, Excel) and interact with them using advanced AI capabilities.
        
        Supported formats:
        - PDF (.pdf)
        - Word Documents (.docx)
        - Text Files (.txt)
        - CSV Files (.csv)
        - Excel Files (.xlsx, .xls)
        """
    )
    
    # Sidebar for file upload and processing
    with st.sidebar:
        st.header("📂 Upload Documents")
        uploaded_files = st.file_uploader(
            "Upload your files here",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt', 'csv', 'xlsx', 'xls']
        )
        
        if st.button("Submit & Process"):
            if uploaded_files:
                with st.spinner("Processing your files..."):
                    raw_text = get_document_text(uploaded_files)
                    if raw_text.strip():  # Check if any text was extracted
                        text_chunks = get_text_chunks(raw_text)
                        vector_store = get_vector_store(text_chunks)
                        if vector_store:
                            st.success("Files processed successfully! You can now ask questions.")
                    else:
                        st.error("No text could be extracted from the uploaded files.")
            else:
                st.error("Please upload files before processing.")

    # Main content area
    st.subheader("Ask Questions")
    user_question = st.text_input("Enter your question about the documents:")
    if user_question:
        with st.spinner("Getting response..."):
            chain = get_conversational_chain()
            if chain:
                vector_store = FAISS.load_local("faiss_index", GoogleGenerativeAIEmbeddings(model="models/embedding-001"), allow_dangerous_deserialization=True)
                response = user_input(user_question, vector_store, chain)
                st.write("### Reply:")
                st.write(response)
                
                # Save question-answer history
                history = st.session_state.get("question_history", [])
                history.append({
                    "question": user_question,
                    "answer": response,
                    "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                })
                st.session_state["question_history"] = history

    # Display question-answer history
    if st.button("Download History"):
        history = st.session_state.get("question_history", [])
        if history:
            history_df = pd.DataFrame(history)
            csv = history_df.to_csv(index=False)
            st.download_button(
                "Download History",
                csv,
                "question_history.csv",
                "text/csv",
                key='download-csv'
            )
        else:
            st.write("No history available.")

if __name__ == "__main__":
    main()